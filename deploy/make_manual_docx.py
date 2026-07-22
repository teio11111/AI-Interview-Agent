# -*- coding: utf-8 -*-
"""生成一键打包脚本使用手册 docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_zh_font(run, size=11, bold=False, color=None):
    """设置中文字体（解决 docx 默认字体对中文不友好的问题）"""
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def add_heading(doc, text, level=1):
    """添加带样式的标题"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 0:
        set_zh_font(run, size=22, bold=True, color=RGBColor(0x4F, 0x46, 0xE5))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        set_zh_font(run, size=16, bold=True, color=RGBColor(0x4F, 0x46, 0xE5))
    elif level == 2:
        set_zh_font(run, size=13, bold=True, color=RGBColor(0x43, 0x38, 0xCA))
    return p


def add_para(doc, text, size=11, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_zh_font(run, size=size, bold=bold, color=color)
    return p


def add_code_block(doc, code_text):
    """添加代码块（灰色背景等宽字体）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # 设置段落背景色
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F1F5F9')
    pPr.append(shd)

    for line in code_text.split('\n'):
        run = p.add_run(line + '\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), 'Consolas')
    return p


def add_table(doc, header, rows):
    """添加简单表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = 'Light Grid Accent 1'
    # 表头
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_zh_font(run, size=11, bold=True)
    # 数据行
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_zh_font(run, size=10)
    return table


def build_docx(out_path):
    doc = Document()

    # 设置全局默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ========== 封面 ==========
    add_heading(doc, 'AI-Interview-Agent', level=0)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('一键打包脚本使用手册')
    set_zh_font(run, size=18, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('make_deploy_zip.ps1')
    set_zh_font(run, size=12, color=RGBColor(0x6B, 0x72, 0x80))

    doc.add_paragraph()
    tip = doc.add_paragraph()
    tip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tip.add_run('适配版本：AI-Interview-Agent v1.0+    |    编写日期：2026-07')
    set_zh_font(run, size=10, color=RGBColor(0x94, 0xA3, 0xB8))

    doc.add_page_break()

    # ========== 章节 0：手册定位 ==========
    add_heading(doc, '手册定位', level=1)
    add_para(doc, '本手册是项目《部署知识点手册》的补充章节，重点讲解本地一键打包脚本 make_deploy_zip.ps1 的使用方法、适用场景与注意事项。')
    add_para(doc, '建议先阅读主手册的"阶段 3：代码打包传输"，了解整体打包-上传-解压-重启的完整链路，再结合本脚本使用。')

    doc.add_page_break()

    # ========== 章节 1：脚本用途 ==========
    add_heading(doc, '一、脚本用途', level=1)
    add_para(doc, '将本地项目一键打包成带时间戳的 zip 文件，自动输出到桌面，并弹出文件夹方便拖给 Xftp，省掉手敲 Compress-Archive 长命令的步骤。')

    add_heading(doc, '二、适用场景', level=1)
    add_para(doc, '· 本地改完代码，准备上传到云端')
    add_para(doc, '· 不想记打包命令的详细参数')
    add_para(doc, '· 想自动区分多次打包的版本（脚本带时间戳）')
    add_para(doc, '· 想要在项目目录也留一份历史打包备份')

    doc.add_page_break()

    # ========== 章节 3：首次部署步骤 ==========
    add_heading(doc, '三、首次部署步骤', level=1)

    add_heading(doc, '步骤 1：解锁 PowerShell 脚本执行策略（仅一次）', level=2)
    add_para(doc, '以管理员身份打开 PowerShell，执行：')
    add_code_block(doc, 'Set-ExecutionPolicy -Scope CurrentUser -ExecutionPolicy RemoteSigned')
    add_para(doc, '输入 Y 回车确认。后续不再需要执行。')

    add_heading(doc, '步骤 2：创建脚本文件', level=2)
    add_para(doc, '在 deploy/ 目录下新建文件 make_deploy_zip.ps1，内容如下：')
    add_code_block(doc, r'''# ============================================
# AI-Interview-Agent 一键打包脚本
# 双击运行：自动生成 ai-interview_时间戳.zip 到桌面
# ============================================

function Write-Step($msg) { Write-Host "`n▶ $msg" -ForegroundColor Cyan }
function Write-OK($msg)    { Write-Host "  ✅ $msg" -ForegroundColor Green }
function Write-Err($msg)   { Write-Host "  ❌ $msg" -ForegroundColor Red }

# 1. 定位项目根目录（脚本在 deploy/ 下，往上一级）
$ScriptDir = Split-Path -Parent $MyInvocation.MyCommand.Path
$ProjectRoot = Split-Path -Parent $ScriptDir
Set-Location $ProjectRoot
Write-OK "项目根目录: $ProjectRoot"

# 2. 生成带时间戳的 zip（避免覆盖旧版本）
$Timestamp = Get-Date -Format "yyyyMMdd_HHmmss"
$ZipName = "ai-interview_$Timestamp.zip"
$DesktopPath = [Environment]::GetFolderPath("Desktop")
$ZipPath = Join-Path $DesktopPath $ZipName
Write-Step "开始打包 → $ZipName"

# 3. 执行打包（排除常见杂物）
try {
    Compress-Archive -Path "$ProjectRoot\*" `
        -DestinationPath $ZipPath `
        -Exclude @(
            "__pycache__",
            ".git",
            "*.log",
            "_*.py",
            ".env",
            "flask.err",
            "flask.log",
            "smoke_*.png",
            "check_cands.png"
        )
    Write-OK "打包完成: $ZipPath"
    Write-OK "文件大小: $((Get-Item $ZipPath).Length / 1KB) KB"
}
catch {
    Write-Err "打包失败: $_"
    pause
    exit 1
}

# 4. 在项目根目录也存一份（备份）
$LocalBackup = Join-Path $ProjectRoot "deploy_lastest.zip"
Copy-Item $ZipPath $LocalBackup -Force
Write-OK "项目内备份: deploy_lastest.zip"

# 5. 提示后续步骤
Write-Step "接下来你要做的："
Write-Host "  1. 打开 Xftp，连上服务器" -ForegroundColor Yellow
Write-Host "  2. 把桌面上的 $ZipName 拖到服务器的 /home/steve/ 目录" -ForegroundColor Yellow
Write-Host "  3. Xshell 跑解压和重启命令（见下方）" -ForegroundColor Yellow

# 6. 自动弹出文件夹
Start-Process explorer.exe "/select,$ZipPath"

Write-Host "`n=========================================" -ForegroundColor Green
Write-Host "  🎉 打包完成！" -ForegroundColor Green
Write-Host "=========================================`n" -ForegroundColor Green

pause''')

    doc.add_page_break()

    # ========== 章节 4：日常使用步骤 ==========
    add_heading(doc, '四、日常使用步骤', level=1)

    add_heading(doc, '步骤 1：本地改完代码', level=2)
    add_para(doc, '在本地项目目录正常修改代码、调试、自测通过。')

    add_heading(doc, '步骤 2：双击运行脚本', level=2)
    add_para(doc, '找到 deploy/make_deploy_zip.ps1 → 右键 → 使用 PowerShell 运行')
    add_para(doc, '期望输出：')
    add_code_block(doc, r'''✅ 项目根目录: C:\Users\Teio\Desktop\AI-Interview-Agent
▶ 开始打包 → ai-interview_20260715_143022.zip
  ✅ 打包完成: C:\Users\Teio\Desktop\ai-interview_20260715_143022.zip
  ✅ 文件大小: 786 KB
  ✅ 项目内备份: deploy_lastest.zip''')

    add_heading(doc, '步骤 3：打开 Xftp 传到服务器', level=2)
    add_para(doc, '把桌面上的 zip 文件拖到服务器的 /home/steve/ 目录。')

    add_heading(doc, '步骤 4：Xshell 里执行 3 行命令', level=2)
    add_para(doc, '在 Xshell 连上服务器后，执行以下命令完成解压、覆盖、重启：')
    add_code_block(doc, r'''unzip -o ~/ai-interview_*.zip -d ~/tmp-deploy
\cp -rf ~/tmp-deploy/* ~/ai-interview/
rm -rf ~/tmp-deploy
bash ~/ai-interview/deploy/restart.sh''')
    add_para(doc, '重启脚本会自动停掉旧的 Flask 进程，再启动新的。整个过程约 5 秒，期间正在进行的面试会话会被中断。')

    doc.add_page_break()

    # ========== 章节 5：注意事项 ==========
    add_heading(doc, '五、注意事项', level=1)
    add_table(doc,
        ['项目', '说明'],
        [
            ['时间戳格式', 'yyyyMMdd_HHmmss，例如 20260715_143022'],
            ['排除的文件', '__pycache__、.git、*.log、.env、测试截图等'],
            ['备份位置', '项目根目录 deploy_lastest.zip（方便溯源）'],
            ['失败处理', '脚本会自动暂停，按任意键退出'],
            ['重复打包', '不会覆盖，每次生成新文件（时间戳不同）'],
            ['依赖 .NET', 'Windows 10/11 自带 Compress-Archive，无需额外安装'],
            ['重启影响', '会中断正在进行的实时面试会话，建议错峰部署'],
        ]
    )

    doc.add_page_break()

    # ========== 章节 6：与原命令对比 ==========
    add_heading(doc, '六、与原打包命令对比', level=1)

    add_heading(doc, '原命令', level=2)
    add_para(doc, '原 UPLOAD_GUIDE.md 里的手动方式：')
    add_code_block(doc, '''Compress-Archive -Path "c:\\Users\\Teio\\Desktop\\AI-Interview-Agent\\*" `
  -DestinationPath "c:\\Users\\Teio\\Desktop\\AI-Interview-Agent\\deploy.zip" `
  -Exclude "__pycache__",".git","*.log","_*.py",".env","flask.err","flask.log","smoke_*.png"''')

    add_heading(doc, '脚本优势', level=2)
    add_table(doc,
        ['维度', '手动命令', '一键脚本'],
        [
            ['上手成本', '需记参数', '双击即用'],
            ['版本区分', '固定文件名 deploy.zip，会覆盖', '带时间戳，保留历史'],
            ['找文件', '手动去项目根目录找', '自动弹出桌面文件夹'],
            ['历史追溯', '无', '项目根目录保留 deploy_lastest.zip'],
            ['部署确认', '需自己比对', '控制台彩色输出+大小提示'],
        ]
    )

    doc.add_page_break()

    # ========== 章节 7：未来升级方向 ==========
    add_heading(doc, '七、未来升级方向（可选）', level=1)
    add_table(doc,
        ['升级', '说明', '工作量'],
        [
            ['自动 SSH 上传', '配 SSH 免密登录后，双击直接上传+重启', '30 分钟'],
            ['自动调起 Xftp', '通过命令行调起 Xftp 的传输协议', '复杂'],
            ['git push 触发', '推到 GitHub/Gitee 后服务器自动拉取', '1-2 小时'],
            ['完整 CI/CD', 'GitHub Actions + rsync deploy key', '2-3 小时'],
        ]
    )

    doc.add_page_break()

    # ========== 章节 8：一句话总结 ==========
    add_heading(doc, '八、一句话总结', level=1)
    p = doc.add_paragraph()
    run = p.add_run('日常部署流程：改代码 → 双击脚本 → 拖 zip 到 Xftp → Xshell 跑 3 行命令 → 完成。整个流程 2 分钟。')
    set_zh_font(run, size=12, bold=True, color=RGBColor(0x4F, 0x46, 0xE5))

    # 保存
    doc.save(out_path)
    print('[OK] Generated:', out_path)


if __name__ == '__main__':
    import os
    out = os.path.join(os.path.dirname(__file__), '一键打包脚本使用手册.docx')
    build_docx(out)